#!/usr/bin/env python3
"""
Professional Resume Matcher with ATS-Grade Parsing
Industry-standard scoring methodology for realistic match percentages.
"""

import streamlit as st
import json
import re
from pathlib import Path
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import pandas as pd
from collections import defaultdict
import io

# ============================================================================
# PDF AND DOCX PARSING - PROFESSIONAL GRADE EXTRACTION
# ============================================================================

def extract_text_from_pdf(file) -> str:
    """Extract text from PDF using professional-grade methods."""
    text = ""
    
    # Method 1: Try pdfplumber (best for tables and formatting)
    try:
        import pdfplumber
        file.seek(0)
        with pdfplumber.open(file) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
                
                # Extract tables (skills often in tables)
                tables = page.extract_tables()
                for table in tables:
                    for row in table:
                        if row:
                            row_text = " | ".join([str(cell) if cell else "" for cell in row])
                            text += row_text + "\n"
        
        if text and len(text.strip()) > 100:
            return clean_extracted_text(text)
    except ImportError:
        pass
    except Exception as e:
        st.warning(f"pdfplumber failed: {e}")
    
    # Fallback: PyPDF2
    try:
        import PyPDF2
        file.seek(0)
        pdf_reader = PyPDF2.PdfReader(file)
        for page in pdf_reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    except Exception as e:
        st.error(f"PDF extraction failed: {e}")
    
    return clean_extracted_text(text)


def extract_text_from_docx(file) -> str:
    """Extract text from DOCX file."""
    text = ""
    try:
        from docx import Document
        file.seek(0)
        doc = Document(file)
        
        # Extract from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text + "\n"
        
        # Extract from tables (often contain skills, education)
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text += " | ".join(row_text) + "\n"
    except Exception as e:
        st.error(f"Error reading DOCX: {e}")
    
    return clean_extracted_text(text)


def clean_extracted_text(text: str) -> str:
    """Clean and normalize extracted text like professional ATS systems."""
    if not text:
        return ""
    
    # Replace common PDF artifacts
    text = text.replace('\x00', '')
    text = re.sub(r'[\x01-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', '', text)
    
    # CRITICAL: Normalize special characters that affect C++, C#, F# detection
    # Common PDF issues: C++ might appear as C＋＋, C＋+, C + +, etc.
    text = re.sub(r'[＋+]{2}', '++', text)  # Normalize plus signs
    text = re.sub(r'c\s*[+＋]{2}', 'c++', text, flags=re.IGNORECASE)
    text = re.sub(r'c\s*#', 'c#', text, flags=re.IGNORECASE)
    text = re.sub(r'f\s*#', 'f#', text, flags=re.IGNORECASE)
    
    # Fix common OCR/extraction issues AFTER special char normalization
    text = re.sub(r'([a-z])([A-Z])', r'\1 \2', text)  # camelCase to spaces
    text = re.sub(r'(\d)([A-Za-z])', r'\1 \2', text)  # 3years -> 3 years
    text = re.sub(r'([A-Za-z])(\d)', r'\1 \2', text)  # Python3 -> Python 3
    
    # Normalize whitespace but preserve newlines for section detection
    text = re.sub(r'[ \t]+', ' ', text)
    text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)
    
    # Fix common ligatures
    replacements = {
        'ﬁ': 'fi', 'ﬂ': 'fl', 'ﬀ': 'ff', 'ﬃ': 'ffi', 'ﬄ': 'ffl',
        '•': ' • ', '●': ' • ', '○': ' ○ ', '■': ' ■ ', '□': ' □ ',
        '–': '-', '—': '-', ''': "'", ''': "'", '"': '"', '"': '"',
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    
    return text.strip()


# ============================================================================
# COMPREHENSIVE SKILL DATABASE (200+ TECHNOLOGIES)
# ============================================================================

# All skills with variations for accurate matching
SKILLS_DATABASE = {
    # Programming Languages - Classic (with ALL common variations)
    "c": [r"\bc\b", r"\bc programming\b", r"\bc language\b", r"\bc lang\b"],
    "c++": [
        r"c\+\+", r"cpp\b", r"c plus plus", r"cplusplus",
        r"c＋＋", r"c \+ \+", r"c\+",  # Common PDF artifacts
    ],
    "c#": [r"c#", r"csharp\b", r"c sharp", r"c＃", r"c-sharp"],
    "java": [r"\bjava\b(?!\s*script)", r"\bjava[,\s]", r"java\s*(?:programming|lang)"],
    "python": [r"\bpython\b", r"python\s*3", r"\bpy\b", r"python[,\s]"],
    "javascript": [r"\bjavascript\b", r"\bjs\b", r"\becmascript\b", r"\bes6\b", r"\bes2015\b"],
    "typescript": [r"\btypescript\b", r"\bts\b"],
    "ruby": [r"\bruby\b"],
    "php": [r"\bphp\b"],
    "perl": [r"\bperl\b"],
    "r": [r"\br programming\b", r"\br language\b", r"\br studio\b", r"\brstats\b"],
    "matlab": [r"\bmatlab\b"],
    "scala": [r"\bscala\b"],
    "kotlin": [r"\bkotlin\b"],
    "swift": [r"\bswift\b"],
    "objective-c": [r"\bobjective-c\b", r"\bobjc\b", r"\bobj-c\b"],
    "go": [r"\bgolang\b", r"\bgo language\b", r"\bgo programming\b"],
    "rust": [r"\brust\b", r"\brustlang\b"],
    "dart": [r"\bdart\b"],
    "lua": [r"\blua\b"],
    "haskell": [r"\bhaskell\b"],
    "clojure": [r"\bclojure\b"],
    "elixir": [r"\belixir\b"],
    "erlang": [r"\berlang\b"],
    "f#": [r"\bf#\b", r"\bfsharp\b"],
    "groovy": [r"\bgroovy\b"],
    "julia": [r"\bjulia\b"],
    "cobol": [r"\bcobol\b"],
    "fortran": [r"\bfortran\b"],
    "pascal": [r"\bpascal\b", r"\bdelphi\b"],
    "assembly": [r"\bassembly\b", r"\basm\b", r"\bx86\b", r"\barm assembly\b"],
    "vba": [r"\bvba\b", r"\bvisual basic\b"],
    "powershell": [r"\bpowershell\b"],
    "bash": [r"\bbash\b", r"\bshell\b", r"\bshell script\b", r"\bzsh\b", r"\bsh\b"],
    "sql": [r"\bsql\b"],
    "plsql": [r"\bpl/sql\b", r"\bplsql\b"],
    "tsql": [r"\bt-sql\b", r"\btsql\b"],
    
    # Web Technologies
    "html": [r"\bhtml\b", r"\bhtml5\b"],
    "css": [r"\bcss\b", r"\bcss3\b"],
    "sass": [r"\bsass\b", r"\bscss\b"],
    "less": [r"\bless\b"],
    "xml": [r"\bxml\b"],
    "json": [r"\bjson\b"],
    "yaml": [r"\byaml\b", r"\byml\b"],
    "graphql": [r"\bgraphql\b"],
    "rest": [r"\brest\b", r"\brestful\b", r"\brest api\b"],
    "soap": [r"\bsoap\b"],
    "ajax": [r"\bajax\b"],
    "websocket": [r"\bwebsocket\b", r"\bwebsockets\b"],
    
    # Frontend Frameworks
    "react": [r"\breact\b", r"\breactjs\b", r"\breact\.js\b", r"\breact native\b"],
    "angular": [r"\bangular\b", r"\bangularjs\b", r"\bangular\.js\b"],
    "vue": [r"\bvue\b", r"\bvuejs\b", r"\bvue\.js\b"],
    "svelte": [r"\bsvelte\b", r"\bsveltekit\b"],
    "next.js": [r"\bnext\.js\b", r"\bnextjs\b", r"\bnext js\b"],
    "nuxt": [r"\bnuxt\b", r"\bnuxtjs\b"],
    "gatsby": [r"\bgatsby\b"],
    "ember": [r"\bember\b", r"\bemberjs\b"],
    "backbone": [r"\bbackbone\b", r"\bbackbonejs\b"],
    "jquery": [r"\bjquery\b"],
    "bootstrap": [r"\bbootstrap\b"],
    "tailwind": [r"\btailwind\b", r"\btailwindcss\b"],
    "material-ui": [r"\bmaterial-ui\b", r"\bmui\b", r"\bmaterial ui\b"],
    "chakra": [r"\bchakra\b", r"\bchakra ui\b"],
    "redux": [r"\bredux\b"],
    "mobx": [r"\bmobx\b"],
    "webpack": [r"\bwebpack\b"],
    "vite": [r"\bvite\b"],
    "parcel": [r"\bparcel\b"],
    "rollup": [r"\brollup\b"],
    "babel": [r"\bbabel\b"],
    
    # Backend Frameworks
    "node.js": [r"\bnode\b", r"\bnodejs\b", r"\bnode\.js\b"],
    "express": [r"\bexpress\b", r"\bexpressjs\b", r"\bexpress\.js\b"],
    "fastify": [r"\bfastify\b"],
    "koa": [r"\bkoa\b"],
    "nestjs": [r"\bnestjs\b", r"\bnest\.js\b"],
    "django": [r"\bdjango\b"],
    "flask": [r"\bflask\b"],
    "fastapi": [r"\bfastapi\b", r"\bfast api\b"],
    "tornado": [r"\btornado\b"],
    "spring": [r"\bspring\b", r"\bspring boot\b", r"\bspringboot\b"],
    "spring mvc": [r"\bspring mvc\b"],
    "hibernate": [r"\bhibernate\b"],
    "rails": [r"\brails\b", r"\bruby on rails\b", r"\bror\b"],
    "sinatra": [r"\bsinatra\b"],
    "laravel": [r"\blaravel\b"],
    "symfony": [r"\bsymfony\b"],
    "codeigniter": [r"\bcodeigniter\b"],
    ".net": [r"\.net\b", r"\bdotnet\b", r"\basp\.net\b"],
    "asp.net core": [r"\basp\.net core\b"],
    "phoenix": [r"\bphoenix\b"],
    "gin": [r"\bgin\b"],
    "echo": [r"\becho\b"],
    "fiber": [r"\bfiber\b"],
    
    # Databases
    "mysql": [r"\bmysql\b"],
    "postgresql": [r"\bpostgresql\b", r"\bpostgres\b", r"\bpsql\b"],
    "sqlite": [r"\bsqlite\b"],
    "oracle": [r"\boracle\b", r"\boracle db\b"],
    "sql server": [r"\bsql server\b", r"\bmssql\b", r"\bms sql\b"],
    "mongodb": [r"\bmongodb\b", r"\bmongo\b"],
    "redis": [r"\bredis\b"],
    "elasticsearch": [r"\belasticsearch\b", r"\belastic search\b"],
    "cassandra": [r"\bcassandra\b"],
    "dynamodb": [r"\bdynamodb\b", r"\bdynamo db\b"],
    "couchdb": [r"\bcouchdb\b"],
    "neo4j": [r"\bneo4j\b"],
    "firebase": [r"\bfirebase\b", r"\bfirestore\b"],
    "supabase": [r"\bsupabase\b"],
    "mariadb": [r"\bmariadb\b"],
    "memcached": [r"\bmemcached\b"],
    "hbase": [r"\bhbase\b"],
    "influxdb": [r"\binfluxdb\b"],
    "timescaledb": [r"\btimescaledb\b"],
    
    # Cloud Platforms
    "aws": [r"\baws\b", r"\bamazon web services\b"],
    "azure": [r"\bazure\b", r"\bmicrosoft azure\b"],
    "gcp": [r"\bgcp\b", r"\bgoogle cloud\b", r"\bgoogle cloud platform\b"],
    "heroku": [r"\bheroku\b"],
    "digitalocean": [r"\bdigitalocean\b", r"\bdigital ocean\b"],
    "vercel": [r"\bvercel\b"],
    "netlify": [r"\bnetlify\b"],
    "cloudflare": [r"\bcloudflare\b"],
    "linode": [r"\blinode\b"],
    "vultr": [r"\bvultr\b"],
    
    # AWS Services
    "ec2": [r"\bec2\b"],
    "s3": [r"\bs3\b", r"\bamazon s3\b"],
    "lambda": [r"\blambda\b", r"\baws lambda\b"],
    "rds": [r"\brds\b"],
    "cloudfront": [r"\bcloudfront\b"],
    "route53": [r"\broute53\b", r"\broute 53\b"],
    "sqs": [r"\bsqs\b"],
    "sns": [r"\bsns\b"],
    "ecs": [r"\becs\b"],
    "eks": [r"\beks\b"],
    "fargate": [r"\bfargate\b"],
    "cloudwatch": [r"\bcloudwatch\b"],
    "iam": [r"\biam\b"],
    
    # DevOps & Tools
    "docker": [r"\bdocker\b", r"\bcontainerization\b"],
    "kubernetes": [r"\bkubernetes\b", r"\bk8s\b"],
    "terraform": [r"\bterraform\b"],
    "ansible": [r"\bansible\b"],
    "puppet": [r"\bpuppet\b"],
    "chef": [r"\bchef\b"],
    "jenkins": [r"\bjenkins\b"],
    "gitlab ci": [r"\bgitlab ci\b", r"\bgitlab-ci\b"],
    "github actions": [r"\bgithub actions\b"],
    "circleci": [r"\bcircleci\b", r"\bcircle ci\b"],
    "travis ci": [r"\btravis\b", r"\btravis ci\b"],
    "nginx": [r"\bnginx\b"],
    "apache": [r"\bapache\b"],
    "linux": [r"\blinux\b", r"\bubuntu\b", r"\bcentos\b", r"\bdebian\b", r"\bred hat\b", r"\bfedora\b"],
    "unix": [r"\bunix\b"],
    "git": [r"\bgit\b(?!hub)(?!lab)"],
    "github": [r"\bgithub\b"],
    "gitlab": [r"\bgitlab\b"],
    "bitbucket": [r"\bbitbucket\b"],
    "svn": [r"\bsvn\b", r"\bsubversion\b"],
    "mercurial": [r"\bmercurial\b", r"\bhg\b"],
    "jira": [r"\bjira\b"],
    "confluence": [r"\bconfluence\b"],
    "trello": [r"\btrello\b"],
    "asana": [r"\basana\b"],
    "slack": [r"\bslack\b"],
    
    # Testing
    "jest": [r"\bjest\b"],
    "mocha": [r"\bmocha\b"],
    "chai": [r"\bchai\b"],
    "jasmine": [r"\bjasmine\b"],
    "cypress": [r"\bcypress\b"],
    "selenium": [r"\bselenium\b"],
    "puppeteer": [r"\bpuppeteer\b"],
    "playwright": [r"\bplaywright\b"],
    "pytest": [r"\bpytest\b"],
    "unittest": [r"\bunittest\b"],
    "junit": [r"\bjunit\b"],
    "testng": [r"\btestng\b"],
    "rspec": [r"\brspec\b"],
    "cucumber": [r"\bcucumber\b"],
    "postman": [r"\bpostman\b"],
    
    # Data Science & ML
    "numpy": [r"\bnumpy\b"],
    "pandas": [r"\bpandas\b"],
    "scipy": [r"\bscipy\b"],
    "matplotlib": [r"\bmatplotlib\b"],
    "seaborn": [r"\bseaborn\b"],
    "plotly": [r"\bplotly\b"],
    "scikit-learn": [r"\bscikit-learn\b", r"\bsklearn\b", r"\bscikit learn\b"],
    "tensorflow": [r"\btensorflow\b", r"\btf\b"],
    "pytorch": [r"\bpytorch\b", r"\btorch\b"],
    "keras": [r"\bkeras\b"],
    "opencv": [r"\bopencv\b", r"\bcv2\b"],
    "nltk": [r"\bnltk\b"],
    "spacy": [r"\bspacy\b"],
    "huggingface": [r"\bhugging face\b", r"\bhuggingface\b", r"\btransformers\b"],
    "langchain": [r"\blangchain\b"],
    "jupyter": [r"\bjupyter\b", r"\bjupyter notebook\b"],
    "anaconda": [r"\banaconda\b", r"\bconda\b"],
    "spark": [r"\bapache spark\b", r"\bspark\b", r"\bpyspark\b"],
    "hadoop": [r"\bhadoop\b", r"\bhdfs\b", r"\bmapreduce\b"],
    "kafka": [r"\bkafka\b", r"\bapache kafka\b"],
    "airflow": [r"\bairflow\b", r"\bapache airflow\b"],
    "mlflow": [r"\bmlflow\b"],
    "dbt": [r"\bdbt\b"],
    "tableau": [r"\btableau\b"],
    "power bi": [r"\bpower bi\b", r"\bpowerbi\b"],
    "looker": [r"\blooker\b"],
    
    # Mobile Development
    "android": [r"\bandroid\b"],
    "ios": [r"\bios\b", r"\biphone\b", r"\bipad\b"],
    "react native": [r"\breact native\b"],
    "flutter": [r"\bflutter\b"],
    "xamarin": [r"\bxamarin\b"],
    "ionic": [r"\bionic\b"],
    "cordova": [r"\bcordova\b", r"\bphonegap\b"],
    "swiftui": [r"\bswiftui\b"],
    "jetpack compose": [r"\bjetpack compose\b"],
    
    # Game Development
    "unity": [r"\bunity\b", r"\bunity3d\b"],
    "unreal": [r"\bunreal\b", r"\bunreal engine\b"],
    "godot": [r"\bgodot\b"],
    "cocos2d": [r"\bcocos2d\b"],
    "phaser": [r"\bphaser\b"],
    
    # Other Tools
    "npm": [r"\bnpm\b"],
    "yarn": [r"\byarn\b"],
    "pip": [r"\bpip\b"],
    "maven": [r"\bmaven\b"],
    "gradle": [r"\bgradle\b"],
    "cmake": [r"\bcmake\b"],
    "make": [r"\bmakefile\b", r"\bgnu make\b"],
    "swagger": [r"\bswagger\b", r"\bopenapi\b"],
    "graphql": [r"\bgraphql\b"],
    "grpc": [r"\bgrpc\b"],
    "rabbitmq": [r"\brabbitmq\b"],
    "celery": [r"\bcelery\b"],
    "socket.io": [r"\bsocket\.io\b", r"\bsocketio\b"],
    
    # Concepts & Methodologies
    "data structures": [r"\bdata structures\b"],
    "algorithms": [r"\balgorithms\b", r"\balgorithm\b"],
    "oop": [r"\boop\b", r"\bobject oriented\b", r"\bobject-oriented\b"],
    "functional programming": [r"\bfunctional programming\b"],
    "design patterns": [r"\bdesign patterns\b"],
    "solid": [r"\bsolid principles\b", r"\bsolid\b"],
    "clean code": [r"\bclean code\b"],
    "tdd": [r"\btdd\b", r"\btest driven\b"],
    "bdd": [r"\bbdd\b", r"\bbehavior driven\b"],
    "ci/cd": [r"\bci/cd\b", r"\bcicd\b", r"\bcontinuous integration\b", r"\bcontinuous deployment\b"],
    "agile": [r"\bagile\b"],
    "scrum": [r"\bscrum\b"],
    "kanban": [r"\bkanban\b"],
    "devops": [r"\bdevops\b"],
    "microservices": [r"\bmicroservices\b", r"\bmicro-services\b"],
    "serverless": [r"\bserverless\b"],
    "soa": [r"\bsoa\b", r"\bservice oriented\b"],
    "mvc": [r"\bmvc\b"],
    "mvvm": [r"\bmvvm\b"],
    "api design": [r"\bapi design\b"],
    "system design": [r"\bsystem design\b"],
    "distributed systems": [r"\bdistributed systems\b"],
    "concurrency": [r"\bconcurrency\b", r"\bmultithreading\b", r"\bparallel\b"],
    "security": [r"\bsecurity\b", r"\bcybersecurity\b", r"\binfosec\b"],
    "cryptography": [r"\bcryptography\b", r"\bencryption\b"],
    "blockchain": [r"\bblockchain\b", r"\bweb3\b", r"\bsolidity\b"],
}

# Skill categories for organized display
SKILL_CATEGORIES = {
    "Programming Languages": [
        "c", "c++", "c#", "java", "python", "javascript", "typescript", "ruby", "php",
        "perl", "r", "matlab", "scala", "kotlin", "swift", "objective-c", "go", "rust",
        "dart", "lua", "haskell", "clojure", "elixir", "erlang", "f#", "groovy", "julia",
        "cobol", "fortran", "pascal", "assembly", "vba", "powershell", "bash", "sql",
        "plsql", "tsql"
    ],
    "Frontend": [
        "html", "css", "sass", "less", "react", "angular", "vue", "svelte", "next.js",
        "nuxt", "gatsby", "ember", "backbone", "jquery", "bootstrap", "tailwind",
        "material-ui", "chakra", "redux", "mobx", "webpack", "vite", "parcel", "rollup", "babel"
    ],
    "Backend": [
        "node.js", "express", "fastify", "koa", "nestjs", "django", "flask", "fastapi",
        "tornado", "spring", "spring mvc", "hibernate", "rails", "sinatra", "laravel",
        "symfony", "codeigniter", ".net", "asp.net core", "phoenix", "gin", "echo", "fiber"
    ],
    "Databases": [
        "mysql", "postgresql", "sqlite", "oracle", "sql server", "mongodb", "redis",
        "elasticsearch", "cassandra", "dynamodb", "couchdb", "neo4j", "firebase",
        "supabase", "mariadb", "memcached", "hbase", "influxdb", "timescaledb"
    ],
    "Cloud & DevOps": [
        "aws", "azure", "gcp", "heroku", "digitalocean", "vercel", "netlify", "cloudflare",
        "docker", "kubernetes", "terraform", "ansible", "puppet", "chef", "jenkins",
        "gitlab ci", "github actions", "circleci", "travis ci", "nginx", "apache",
        "linux", "unix", "git", "github", "gitlab", "bitbucket"
    ],
    "Data Science & ML": [
        "numpy", "pandas", "scipy", "matplotlib", "seaborn", "plotly", "scikit-learn",
        "tensorflow", "pytorch", "keras", "opencv", "nltk", "spacy", "huggingface",
        "langchain", "jupyter", "anaconda", "spark", "hadoop", "kafka", "airflow",
        "mlflow", "dbt", "tableau", "power bi", "looker"
    ],
    "Mobile": [
        "android", "ios", "react native", "flutter", "xamarin", "ionic", "cordova",
        "swiftui", "jetpack compose"
    ],
    "Testing": [
        "jest", "mocha", "chai", "jasmine", "cypress", "selenium", "puppeteer",
        "playwright", "pytest", "unittest", "junit", "testng", "rspec", "cucumber", "postman"
    ],
}

# Related skills for semantic matching
RELATED_SKILLS = {
    "python": ["django", "flask", "fastapi", "pandas", "numpy", "pytorch", "tensorflow", "scikit-learn"],
    "javascript": ["react", "angular", "vue", "node.js", "express", "next.js", "typescript"],
    "typescript": ["react", "angular", "node.js", "next.js"],
    "java": ["spring", "hibernate", "kotlin", "android", "maven", "gradle"],
    "c++": ["c", "cmake", "qt"],
    "c#": [".net", "asp.net core", "unity", "xamarin"],
    "ruby": ["rails", "sinatra", "rspec"],
    "php": ["laravel", "symfony", "codeigniter"],
    "go": ["docker", "kubernetes", "gin"],
    "react": ["redux", "next.js", "javascript", "typescript"],
    "angular": ["typescript", "rxjs"],
    "vue": ["nuxt", "javascript"],
    "node.js": ["express", "javascript", "npm"],
    "django": ["python", "postgresql"],
    "spring": ["java", "hibernate", "maven"],
    "aws": ["ec2", "s3", "lambda", "docker", "kubernetes"],
    "docker": ["kubernetes", "linux", "ci/cd"],
    "kubernetes": ["docker", "helm", "aws", "gcp", "azure"],
    "machine learning": ["python", "tensorflow", "pytorch", "scikit-learn", "pandas", "numpy"],
    "data science": ["python", "pandas", "numpy", "sql", "tableau", "jupyter"],
}

# Top universities
TOP_UNIVERSITIES = [
    "stanford", "mit", "massachusetts institute of technology", "berkeley", "uc berkeley",
    "carnegie mellon", "cmu", "harvard", "princeton", "yale", "columbia", "cornell",
    "caltech", "california institute of technology", "georgia tech", "georgia institute",
    "university of washington", "uw seattle", "ucla", "usc", "university of southern california",
    "university of michigan", "umich", "university of illinois", "uiuc", "purdue",
    "university of texas", "ut austin", "university of pennsylvania", "upenn", "duke",
    "northwestern", "brown", "dartmouth", "rice", "vanderbilt", "notre dame", "nyu",
    "university of wisconsin", "ohio state", "penn state", "university of maryland",
    "virginia tech", "university of colorado", "uc san diego", "ucsd", "uc davis",
    "uc irvine", "uc santa barbara", "boston university", "northeastern"
]

# Top tech companies
TOP_COMPANIES = [
    "google", "meta", "facebook", "amazon", "apple", "microsoft", "netflix", "nvidia",
    "stripe", "airbnb", "uber", "lyft", "doordash", "coinbase", "robinhood", "square",
    "databricks", "snowflake", "figma", "notion", "discord", "slack", "dropbox",
    "twitter", "x corp", "linkedin", "salesforce", "adobe", "intel", "amd", "qualcomm",
    "openai", "anthropic", "deepmind", "palantir", "snap", "pinterest", "tiktok", "bytedance",
    "spotify", "oracle", "ibm", "cisco", "vmware", "intuit", "paypal", "visa", "mastercard",
    "goldman sachs", "morgan stanley", "jpmorgan", "jane street", "citadel", "two sigma",
    "bloomberg", "tesla", "spacex", "waymo", "cruise", "rivian", "lucid"
]


# ============================================================================
# SKILL EXTRACTION
# ============================================================================

def extract_skills(text: str) -> tuple:
    """Extract skills from text using multiple methods like professional ATS."""
    text_lower = text.lower()
    found_skills = set()
    skills_by_category = defaultdict(list)
    skill_contexts = {}  # Track where skills were found
    
    # Method 1: Pattern matching with context
    for skill_name, patterns in SKILLS_DATABASE.items():
        skill_found = False
        for pattern in patterns:
            try:
                matches = list(re.finditer(pattern, text_lower, re.IGNORECASE))
                if matches:
                    found_skills.add(skill_name)
                    # Store context for debugging
                    for match in matches[:1]:
                        start = max(0, match.start() - 25)
                        end = min(len(text), match.end() + 25)
                        skill_contexts[skill_name] = text[start:end].strip()
                    
                    # Find category
                    for category, skills in SKILL_CATEGORIES.items():
                        if skill_name in skills:
                            if skill_name not in skills_by_category[category]:
                                skills_by_category[category].append(skill_name)
                            break
                    skill_found = True
                    break
            except re.error:
                continue
    
    # Method 2: Exact word matching for problematic skills (C++, C#, etc.)
    # This catches cases where regex fails due to special characters
    exact_matches = {
        'c++': ['c++', 'cpp', 'c plus plus', 'cplusplus'],
        'c#': ['c#', 'csharp', 'c sharp', 'c-sharp'],
        'f#': ['f#', 'fsharp', 'f sharp', 'f-sharp'],
    }
    
    for skill_name, variants in exact_matches.items():
        if skill_name not in found_skills:
            for variant in variants:
                if variant in text_lower:
                    found_skills.add(skill_name)
                    skill_contexts[skill_name] = f"Found: '{variant}'"
                    # Add to category
                    for category, skills in SKILL_CATEGORIES.items():
                        if skill_name in skills:
                            if skill_name not in skills_by_category[category]:
                                skills_by_category[category].append(skill_name)
                            break
                    break
    
    return found_skills, dict(skills_by_category)


# ============================================================================
# EDUCATION EXTRACTION
# ============================================================================

def extract_education(text: str) -> dict:
    """Extract education details from resume."""
    text_lower = text.lower()
    
    education = {
        "degree_level": "unknown",
        "degree_score": 0.5,
        "gpa": None,
        "major": None,
        "university": None,
        "is_top_university": False,
        "is_cs_related": False,
        "graduation_year": None
    }
    
    # Degree level detection with scores
    degree_patterns = [
        (r"\bph\.?d\.?\b|\bdoctorate\b|\bdoctoral\b", "PhD", 1.0),
        (r"\bmaster'?s?\b|\bm\.?s\.?\b|\bm\.?eng\.?\b|\bmba\b", "Master's", 0.9),
        (r"\bbachelor'?s?\b|\bb\.?s\.?\b|\bb\.?a\.?\b|\bb\.?eng\.?\b|\bundergraduate\b", "Bachelor's", 0.8),
        (r"\bassociate'?s?\b|\ba\.?s\.?\b|\ba\.?a\.?\b", "Associate's", 0.6),
        (r"\bstudent\b|\bpursuing\b|\bcurrently enrolled\b|\bexpected\b", "Student", 0.7),
    ]
    
    for pattern, degree, score in degree_patterns:
        if re.search(pattern, text_lower):
            if score > education["degree_score"]:
                education["degree_level"] = degree
                education["degree_score"] = score
    
    # GPA extraction - multiple patterns
    gpa_patterns = [
        r"gpa[:\s]*([0-4]\.[0-9]{1,2})",
        r"([0-4]\.[0-9]{1,2})\s*/\s*4\.0",
        r"([0-4]\.[0-9]{1,2})\s*gpa",
        r"cumulative[:\s]*([0-4]\.[0-9]{1,2})",
        r"grade point[:\s]*([0-4]\.[0-9]{1,2})",
    ]
    
    for pattern in gpa_patterns:
        match = re.search(pattern, text_lower)
        if match:
            try:
                gpa = float(match.group(1))
                if 0.0 <= gpa <= 4.0:
                    education["gpa"] = gpa
                    break
            except ValueError:
                continue
    
    # CS-related major detection
    cs_majors = [
        r"computer science", r"computer engineering", r"software engineering",
        r"electrical engineering", r"information technology", r"data science",
        r"information systems", r"computational", r"artificial intelligence",
        r"machine learning", r"cybersecurity", r"computer information",
        r"\bcs\b", r"\bece\b", r"\beecs\b", r"\bcse\b"
    ]
    
    for pattern in cs_majors:
        if re.search(pattern, text_lower):
            education["is_cs_related"] = True
            education["major"] = pattern.replace(r"\b", "").replace("\\", "")
            break
    
    # Top university detection
    for uni in TOP_UNIVERSITIES:
        if uni in text_lower:
            education["is_top_university"] = True
            education["university"] = uni.title()
            break
    
    # Graduation year
    year_patterns = [
        r"class of[:\s]*20(\d{2})",
        r"expected[:\s]*(?:graduation)?[:\s]*(?:may|june|december|spring|fall|winter|summer)?[,\s]*20(\d{2})",
        r"graduat(?:ed?|ing|ion)[:\s]*(?:in)?[:\s]*(?:may|june|december|spring|fall|winter|summer)?[,\s]*20(\d{2})",
        r"20(\d{2})\s*[-–]\s*(?:present|current|expected)",
    ]
    
    for pattern in year_patterns:
        match = re.search(pattern, text_lower)
        if match:
            try:
                year = int("20" + match.group(1))
                if 2015 <= year <= 2030:
                    education["graduation_year"] = year
                    break
            except ValueError:
                continue
    
    return education


# ============================================================================
# EXPERIENCE EXTRACTION
# ============================================================================

def extract_experience(text: str) -> dict:
    """Extract work and project experience."""
    text_lower = text.lower()
    
    experience = {
        "has_internship": False,
        "internship_count": 0,
        "has_full_time": False,
        "has_top_company": False,
        "top_companies_worked": [],
        "project_count": 0,
        "has_research": False,
        "has_publications": False,
        "has_leadership": False,
        "years_experience": 0,
    }
    
    # Internship detection
    internship_patterns = [
        r"\bintern\b", r"\binternship\b", r"\bco-?op\b", r"\bsummer\s+\d{4}\b",
        r"\bfall\s+\d{4}\b", r"\bwinter\s+\d{4}\b", r"\bspring\s+\d{4}\b"
    ]
    
    for pattern in internship_patterns:
        matches = re.findall(pattern, text_lower)
        if matches:
            experience["has_internship"] = True
            experience["internship_count"] = max(experience["internship_count"], len(matches))
    
    # Full-time experience
    if re.search(r"\bfull[- ]?time\b|\bemployee\b|\bstaff\b", text_lower):
        experience["has_full_time"] = True
    
    # Top company detection
    for company in TOP_COMPANIES:
        if company in text_lower:
            experience["has_top_company"] = True
            if company.title() not in experience["top_companies_worked"]:
                experience["top_companies_worked"].append(company.title())
    
    # Project counting
    project_indicators = [
        r"\bproject[s]?\b", r"\bbuilt\b", r"\bdeveloped\b", r"\bcreated\b",
        r"\bimplemented\b", r"\bdesigned\b", r"\barchitected\b"
    ]
    
    project_count = 0
    for pattern in project_indicators:
        matches = re.findall(pattern, text_lower)
        project_count += len(matches)
    experience["project_count"] = min(project_count // 2, 15)  # Normalize
    
    # Research and publications
    if re.search(r"\bresearch\b|\bresearcher\b|\blab\b|\blaboratory\b", text_lower):
        experience["has_research"] = True
    
    if re.search(r"\bpublication\b|\bpublished\b|\bpaper\b|\bjournal\b|\bconference\b|\barxiv\b", text_lower):
        experience["has_publications"] = True
    
    # Leadership
    leadership_patterns = [
        r"\bled\b", r"\bleader\b", r"\bleadership\b", r"\bmanaged\b", r"\bmanager\b",
        r"\bpresident\b", r"\bvice president\b", r"\bfounder\b", r"\bco-?founder\b",
        r"\bcaptain\b", r"\bhead\b", r"\bdirector\b", r"\bcoordinator\b", r"\borganized\b"
    ]
    
    for pattern in leadership_patterns:
        if re.search(pattern, text_lower):
            experience["has_leadership"] = True
            break
    
    # Years of experience
    year_exp_patterns = [
        r"(\d+)\+?\s*years?\s*(?:of\s*)?(?:experience|exp)",
        r"experience[:\s]*(\d+)\+?\s*years?",
    ]
    
    for pattern in year_exp_patterns:
        match = re.search(pattern, text_lower)
        if match:
            try:
                years = int(match.group(1))
                experience["years_experience"] = min(years, 15)
                break
            except ValueError:
                continue
    
    return experience


# ============================================================================
# RESUME ANALYSIS
# ============================================================================

def analyze_resume(text: str) -> dict:
    """Complete resume analysis."""
    skills, skills_by_category = extract_skills(text)
    education = extract_education(text)
    experience = extract_experience(text)
    
    return {
        "skills": skills,
        "skills_by_category": skills_by_category,
        "education": education,
        "experience": experience,
        "word_count": len(text.split()),
        "text": text
    }


# ============================================================================
# MATCHING ALGORITHM
# ============================================================================

def calculate_match(resume: dict, job: dict) -> dict:
    """
    Calculate realistic ATS-style match score between resume and job.
    Uses industry-standard scoring with conservative percentages.
    """
    job_title = job.get("title", "").lower()
    job_desc = job.get("description", "").lower()
    company = job.get("company", "")
    job_text = f"{job_title} {job_desc}"
    
    # Extract job requirements
    job_skills, _ = extract_skills(job_text)
    
    # Baseline SWE intern skills if job description is generic
    baseline_swe_skills = {"python", "java", "c++", "javascript", "data structures", "algorithms", "git", "sql"}
    if len(job_skills) < 3:
        job_skills = baseline_swe_skills
    
    resume_skills = resume["skills"]
    
    # ========================================
    # 1. REQUIRED SKILLS MATCHING (50%)
    # ========================================
    # This is the most critical factor for ATS systems
    
    # Core requirements (must-haves)
    core_requirements = job_skills & baseline_swe_skills
    if not core_requirements:
        core_requirements = {"python", "java", "data structures", "algorithms"}
    
    direct_matches = resume_skills & job_skills
    core_matches = resume_skills & core_requirements
    
    # Penalize heavily for missing core skills
    if core_requirements:
        core_match_ratio = len(core_matches) / len(core_requirements)
    else:
        core_match_ratio = 0.7
    
    # Overall skill match
    if job_skills:
        skill_match_ratio = len(direct_matches) / len(job_skills)
    else:
        skill_match_ratio = 0.5
    
    # Related skills (partial credit)
    related_matches = set()
    for skill in resume_skills:
        if skill in RELATED_SKILLS:
            for related in RELATED_SKILLS[skill]:
                if related in job_skills and related not in direct_matches:
                    related_matches.add(related)
    
    related_bonus = min(len(related_matches) * 0.03, 0.15)
    
    # CONSERVATIVE skill scoring
    skill_score = (
        core_match_ratio * 0.50 +      # 50% weight on core skills
        skill_match_ratio * 0.40 +      # 40% weight on all skills
        related_bonus                    # Up to 15% for related skills
    )
    
    # No bonus for skill quantity - only what matches matters
    skill_score = min(skill_score, 1.0)
    
    # ========================================
    # 2. EXPERIENCE QUALITY (30%)
    # ========================================
    # Realistic: most interns have LIMITED experience
    
    exp = resume["experience"]
    exp_score = 0.3  # Lower baseline (most students are early career)
    
    # Internship experience (realistic weight)
    if exp["has_internship"]:
        exp_score += 0.20
        if exp["internship_count"] >= 2:
            exp_score += 0.15
        elif exp["internship_count"] >= 3:
            exp_score += 0.10  # Diminishing returns
    
    # FAANG/top company (significant but not overwhelming)
    if exp["has_top_company"]:
        exp_score += 0.20
        # Same company bonus
        for top_co in exp["top_companies_worked"]:
            if top_co.lower() in company.lower():
                exp_score += 0.05
                break
    
    # Projects (moderate weight)
    if exp["project_count"] >= 8:
        exp_score += 0.10
    elif exp["project_count"] >= 5:
        exp_score += 0.07
    elif exp["project_count"] >= 3:
        exp_score += 0.04
    
    # Research/publications (niche benefit)
    if exp["has_research"]:
        if any(kw in job_title for kw in ["research", "ml", "machine learning", "ai", "data", "scientist"]):
            exp_score += 0.08
        else:
            exp_score += 0.03
    
    if exp["has_publications"]:
        exp_score += 0.05
    
    # Leadership (small bonus)
    if exp["has_leadership"]:
        exp_score += 0.03
    
    exp_score = min(exp_score, 1.0)
    
    # ========================================
    # 3. EDUCATION QUALIFICATIONS (15%)
    # ========================================
    # Important but not decisive for interns
    
    edu = resume["education"]
    edu_score = edu["degree_score"] * 0.5  # Reduce base impact
    
    # CS major (essential for SWE)
    if edu["is_cs_related"]:
        edu_score += 0.30
    else:
        edu_score += 0.10  # Small credit for any degree
    
    # GPA (matters but not everything)
    if edu["gpa"]:
        if edu["gpa"] >= 3.9:
            edu_score += 0.15
        elif edu["gpa"] >= 3.7:
            edu_score += 0.10
        elif edu["gpa"] >= 3.5:
            edu_score += 0.05
        elif edu["gpa"] >= 3.0:
            edu_score += 0.02
    
    # Top university (competitive advantage)
    if edu["is_top_university"]:
        edu_score += 0.08
    
    edu_score = min(edu_score, 1.0)
    
    # ========================================
    # 4. SEMANTIC FIT (5%)
    # ========================================
    # Least important - skills and experience matter more
    
    try:
        vectorizer = TfidfVectorizer(
            stop_words='english',
            ngram_range=(1, 2),
            max_features=500,
            min_df=1
        )
        docs = [resume["text"][:10000], job_text[:5000]]  # Limit size
        tfidf = vectorizer.fit_transform(docs)
        text_score = cosine_similarity(tfidf[0:1], tfidf[1:2])[0][0]
    except:
        text_score = 0.2
    
    # ========================================
    # FINAL REALISTIC SCORE
    # ========================================
    
    final_score = (
        skill_score * 0.50 +    # Skills are king for ATS
        exp_score * 0.30 +      # Experience is important
        edu_score * 0.15 +      # Education matters moderately
        text_score * 0.05       # Semantic fit is minor
    ) * 100
    
    # Apply reality check penalties
    
    # Penalty for missing core skills
    if core_match_ratio < 0.5:
        final_score *= 0.75  # 25% penalty
    elif core_match_ratio < 0.75:
        final_score *= 0.90  # 10% penalty
    
    # Penalty for having few total skills
    if len(resume_skills) < 5:
        final_score *= 0.80
    elif len(resume_skills) < 10:
        final_score *= 0.90
    
    # Realistic floor and ceiling
    # Very few candidates score >80% realistically
    final_score = max(15, min(88, final_score))
    
    # Additional ceiling for lacking experience
    if not exp["has_internship"] and not exp["has_top_company"]:
        final_score = min(final_score, 65)
    
    # Quality labels (more conservative)
    if final_score >= 75:
        quality = "⭐ Strong Match"
    elif final_score >= 60:
        quality = "✓ Good Match"
    elif final_score >= 45:
        quality = "~ Fair Match"
    elif final_score >= 30:
        quality = "⚠ Weak Match"
    else:
        quality = "✗ Poor Match"
    
    # Missing skills
    missing = list(job_skills - resume_skills - related_matches)[:5]
    
    return {
        "score": round(final_score, 1),
        "quality": quality,
        "breakdown": {
            "Skills": round(skill_score * 100, 1),
            "Experience": round(exp_score * 100, 1),
            "Education": round(edu_score * 100, 1),
            "Relevance": round(text_score * 100, 1)
        },
        "matched_skills": list(direct_matches)[:10],
        "related_skills": list(related_matches)[:5],
        "missing_skills": missing
    }


# ============================================================================
# DATA LOADING
# ============================================================================

def load_jobs() -> list:
    """Load jobs from jobs.json."""
    jobs_file = Path("jobs.json")
    if not jobs_file.exists():
        return []
    
    try:
        with open(jobs_file) as f:
            data = json.load(f)
            if isinstance(data, dict):
                jobs = []
                for category, job_list in data.items():
                    for job in job_list:
                        job["category"] = "FAANG+" if category.lower() == "faang" else category.title()
                        jobs.append(job)
                return jobs
            return data
    except Exception as e:
        return []


# ============================================================================
# STREAMLIT APP
# ============================================================================

def main():
    st.set_page_config(
        page_title="Resume Matcher",
        page_icon="🎯",
        layout="wide"
    )
    
    st.title("🎯 AI Resume Matcher")
    st.markdown("Upload your resume to get **accurate match percentages** for SWE internship positions.")
    
    # Sidebar
    with st.sidebar:
        st.header("📄 Upload Resume")
        
        uploaded_file = st.file_uploader(
            "Choose file",
            type=["pdf", "docx", "txt"],
            help="PDF, DOCX, or TXT"
        )
        
        resume_analysis = None
        
        if uploaded_file:
            file_type = uploaded_file.name.split(".")[-1].lower()
            
            with st.spinner("Parsing resume..."):
                if file_type == "pdf":
                    text = extract_text_from_pdf(uploaded_file)
                elif file_type == "docx":
                    text = extract_text_from_docx(uploaded_file)
                else:
                    text = uploaded_file.read().decode("utf-8")
            
            if text and len(text.strip()) > 50:
                resume_analysis = analyze_resume(text)
                st.success(f"✅ Parsed {resume_analysis['word_count']} words")
                
                # DEBUG: Show raw extracted text preview
                with st.expander("🔍 Debug: View Extracted Text"):
                    st.text_area("Raw Text (first 1500 chars)", text[:1500], height=200)
                    st.caption("Check if your skills appear here. If C++ is missing, the PDF extraction failed.")
                
                # Show parsed data
                st.markdown("---")
                
                # Skills
                st.markdown("### 🔧 Skills Found")
                total_skills = len(resume_analysis["skills"])
                st.markdown(f"**{total_skills} skills detected**")
                
                for category, skills in resume_analysis["skills_by_category"].items():
                    with st.expander(f"{category} ({len(skills)})"):
                        st.write(", ".join(skills))
                
                # Education
                st.markdown("---")
                st.markdown("### 🎓 Education")
                edu = resume_analysis["education"]
                st.write(f"**Degree:** {edu['degree_level']}")
                if edu["gpa"]:
                    st.write(f"**GPA:** {edu['gpa']}")
                if edu["is_cs_related"]:
                    st.write("✅ CS/Tech Major")
                if edu["is_top_university"]:
                    st.write(f"✅ Top University: {edu['university']}")
                
                # Experience
                st.markdown("---")
                st.markdown("### 💼 Experience")
                exp = resume_analysis["experience"]
                if exp["has_internship"]:
                    st.write(f"✅ {exp['internship_count']} internship(s)")
                if exp["has_top_company"]:
                    st.write(f"✅ Top companies: {', '.join(exp['top_companies_worked'][:3])}")
                st.write(f"📁 ~{exp['project_count']} projects")
                if exp["has_research"]:
                    st.write("🔬 Research experience")
                if exp["has_leadership"]:
                    st.write("👤 Leadership experience")
            else:
                st.error("Could not extract text. Try a different file format.")
        
        st.markdown("---")
        st.markdown("### Filters")
        min_match = st.slider("Min Match %", 0, 100, 0, 5)
        category_filter = st.selectbox("Category", ["All", "FAANG+", "Other"])
    
    # Main content
    jobs = load_jobs()
    
    if not jobs:
        st.warning("No jobs found. Run `python scraper.py` first.")
        return
    
    if not resume_analysis:
        st.info("👈 Upload your resume to see match scores!")
        st.subheader(f"📋 {len(jobs)} Available Positions")
        
        df = pd.DataFrame([{
            "Company": j.get("company", ""),
            "Role": j.get("title", ""),
            "Location": j.get("location", ""),
        } for j in jobs])
        st.dataframe(df, use_container_width=True, hide_index=True)
        return
    
    # Calculate matches
    st.subheader("🎯 Your Matches")
    
    results = []
    progress = st.progress(0)
    
    for i, job in enumerate(jobs):
        match = calculate_match(resume_analysis, job)
        results.append({
            "company": job.get("company", ""),
            "title": job.get("title", ""),
            "location": job.get("location", ""),
            "category": job.get("category", "Other"),
            "url": job.get("url", job.get("apply_url", "")),
            **match
        })
        progress.progress((i + 1) / len(jobs))
    
    progress.empty()
    
    # Sort and filter
    results.sort(key=lambda x: x["score"], reverse=True)
    
    if min_match > 0:
        results = [r for r in results if r["score"] >= min_match]
    
    if category_filter != "All":
        results = [r for r in results if category_filter.replace("+", "") in r["category"]]
    
    if not results:
        st.warning("No matches found with current filters.")
        return
    
    # Stats
    col1, col2, col3, col4 = st.columns(4)
    col1.metric("Total", len(results))
    col2.metric("Average", f"{sum(r['score'] for r in results) / len(results):.1f}%")
    col3.metric("Excellent (≥75%)", len([r for r in results if r["score"] >= 75]))
    col4.metric("Best", f"{results[0]['score']}%")
    
    st.markdown("---")
    
    # Results
    for i, r in enumerate(results):
        icon = "🟢" if r["score"] >= 75 else "🟡" if r["score"] >= 50 else "🔴"
        
        with st.expander(f"{icon} **{r['score']}%** | {r['company']} - {r['title']}", expanded=(i < 3)):
            c1, c2 = st.columns([3, 1])
            
            with c1:
                st.markdown(f"**{r['company']}** — {r['title']}")
                st.markdown(f"📍 {r['location']} | {r['quality']}")
                
                # Breakdown
                st.markdown("**Score Breakdown:**")
                cols = st.columns(4)
                for idx, (k, v) in enumerate(r["breakdown"].items()):
                    cols[idx].metric(k, f"{v}%")
                
                # Skills
                if r["matched_skills"]:
                    st.markdown(f"✅ **Matched:** {', '.join(r['matched_skills'])}")
                if r["related_skills"]:
                    st.markdown(f"🔗 **Related:** {', '.join(r['related_skills'])}")
                if r["missing_skills"]:
                    st.markdown(f"📝 **Consider adding:** {', '.join(r['missing_skills'])}")
            
            with c2:
                if r["url"]:
                    st.link_button("🔗 Apply", r["url"], use_container_width=True)
    
    # Export
    st.markdown("---")
    df = pd.DataFrame([{
        "Company": r["company"],
        "Role": r["title"],
        "Location": r["location"],
        "Match %": r["score"],
        "Quality": r["quality"]
    } for r in results])
    
    st.download_button("📥 Download CSV", df.to_csv(index=False), "matches.csv", "text/csv")


if __name__ == "__main__":
    main()
